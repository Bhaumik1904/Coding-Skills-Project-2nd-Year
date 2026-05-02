from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph border bottom ───────────────────────────────────────
def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1B4F72')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1B, 0x4F, 0x72)
MID_BLUE    = RGBColor(0x21, 0x80, 0xC8)
ACCENT_GOLD = RGBColor(0xE6, 0x7E, 0x22)
LIGHT_GREY  = RGBColor(0xF2, 0xF3, 0xF4)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT   = RGBColor(0x17, 0x20, 0x2A)

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Logo
logo_path = r"d:\Fractional Knapsack\screenshots\srm_logo.png"
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(1.8))

doc.add_paragraph()

# University name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SRM UNIVERSITY AP, AMARAVATI")
run.font.size  = Pt(16)
run.font.bold  = True
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Department of Computer Science and Engineering")
run.font.size  = Pt(12)
run.font.color.rgb = DARK_BLUE

# Divider
p = doc.add_paragraph("─" * 60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.color.rgb = ACCENT_GOLD

doc.add_paragraph()

# Main title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FRACTIONAL KNAPSACK VISUALIZER")
run.font.size  = Pt(24)
run.font.bold  = True
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Interactive Algorithm Comparison Tool")
run.font.size  = Pt(14)
run.font.italic = True
run.font.color.rgb = MID_BLUE

doc.add_paragraph()

# Badge line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("■  PROJECT REPORT  ■")
run.font.size  = Pt(13)
run.font.bold  = True
run.font.color.rgb = ACCENT_GOLD

doc.add_paragraph()

# Course info table (borderless)
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Subject",          "Coding Skills"),
    ("Subject Code",     "CSE 1101"),
    ("Academic Year",    "2025 – 2026"),
    ("Year / Semester",  "2nd Year – 4th Semester"),
    ("Section",          "C"),
]
for i, (label, value) in enumerate(info_data):
    row = info_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for j, cell in enumerate(row.cells):
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(11)
        if j == 0:
            run.bold = True
            run.font.color.rgb = DARK_BLUE
        else:
            run.font.color.rgb = DARK_TEXT
    # borders
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBdr = OxmlElement('w:tcBdr')
        for side in ('top','left','bottom','right'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   'none')
            el.set(qn('w:sz'),    '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
            tcBdr.append(el)
        tcPr.append(tcBdr)

doc.add_paragraph()

# Submitted by
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted by")
run.font.size  = Pt(11)
run.font.italic = True
run.font.color.rgb = DARK_TEXT

doc.add_paragraph()

# Team table
team = [
    ("Bhaumik Hinunia",        "AP24110010182"),
    ("Manya Srivastava",       "AP24110010171"),
    ("Vardhana Paluvai",       "AP24110012136"),
    ("Jayanti Yadav",          "AP24110010132"),
    ("Vijithendra Nagabhyru",  "AP24110010184"),
]

team_table = doc.add_table(rows=len(team)+1, cols=2)
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
team_table.style = 'Table Grid'

# Header row
hdr_cells = team_table.rows[0].cells
hdr_cells[0].text = "Name"
hdr_cells[1].text = "Enrollment Number"
for cell in hdr_cells:
    set_cell_bg(cell, "1B4F72")
    run = cell.paragraphs[0].runs[0]
    run.font.bold  = True
    run.font.color.rgb = WHITE
    run.font.size  = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (name, enroll) in enumerate(team):
    row = team_table.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = enroll
    bg = "EBF5FB" if i % 2 == 0 else "FFFFFF"
    for cell in row.cells:
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_TEXT
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("May 2026")
run.font.size  = Pt(11)
run.font.color.rgb = DARK_TEXT

# ── Page Break ────────────────────────────────────────────────────────────────
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  HELPER: SECTION HEADING
# ══════════════════════════════════════════════════════════════════════════════
def section_heading(number, title):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}.  {title.upper()}")
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE
    add_bottom_border(p)
    doc.add_paragraph()

def body(text, size=11):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = DARK_TEXT
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after = Pt(3)
    return p

# ══════════════════════════════════════════════════════════════════════════════
#  1. ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
section_heading("1", "Abstract")
body(
    "This report presents the Fractional Knapsack Visualizer, an interactive "
    "web-based educational tool developed as the Final Project for the Coding Skills "
    "course (2nd Year, 4th Semester) at SRM University AP, Amaravati. The application "
    "visually demonstrates and compares the classic Greedy Algorithm against a Naive "
    "Sequential Selection approach for solving the Fractional Knapsack problem. "
    "Built with React 19 and TypeScript, the tool provides step-by-step animated "
    "simulations, real-time profit statistics, and interactive charts — enabling "
    "students to intuitively grasp algorithm efficiency and optimization concepts."
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  2. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
section_heading("2", "Introduction")
body(
    "The Fractional Knapsack problem is a classic optimization problem in computer science "
    "and operations research. Given a set of items, each with a weight and profit, the goal "
    "is to determine the maximum-value subset of items that fit within a given weight capacity — "
    "where items can be split into fractions."
)
body(
    "While the theoretical solution is well-established, visual understanding often remains "
    "a challenge for learners. This project bridges that gap through an interactive simulation "
    "that walks users through each step of both the Greedy and Naive approaches in real time, "
    "using a cargo-loading metaphor (a ship with limited tonnage capacity) to make the "
    "concept relatable and easy to grasp."
)
body(
    "The project is fully deployed and accessible at:"
)
p = doc.add_paragraph()
run = p.add_run("🔗  https://coding-skills-project-2nd-year.vercel.app/")
run.font.size  = Pt(11)
run.font.color.rgb = MID_BLUE
run.font.bold  = True
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  3. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
section_heading("3", "Problem Statement")
body(
    "Given a knapsack (or ship) of capacity W and n items each having a weight wᵢ and profit pᵢ, "
    "select the amounts xᵢ (0 ≤ xᵢ ≤ wᵢ) such that:"
)
p = doc.add_paragraph()
run = p.add_run("    Maximize  ∑ pᵢ · (xᵢ / wᵢ)    subject to  ∑ xᵢ ≤ W")
run.font.size  = Pt(11)
run.font.bold  = True
run.font.color.rgb = DARK_BLUE
body(
    "This project visualizes two approaches to solving this problem: "
    "the optimal Greedy Algorithm and a suboptimal Naive Sequential approach, "
    "clearly demonstrating the superiority of the greedy strategy."
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  4. ALGORITHM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
section_heading("4", "Algorithm Overview")

p = doc.add_paragraph()
run = p.add_run("4.1  Greedy Algorithm (Optimal)")
run.font.size  = Pt(12)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x1A, 0x8A, 0x44)
doc.add_paragraph()

bullet("Compute the profit-to-weight ratio (pᵢ / wᵢ) for every item.")
bullet("Sort all items in descending order of their ratio.")
bullet("Greedily pick items one by one: if the full item fits, take it entirely; otherwise take a fractional portion to fill the remaining capacity.")
bullet("Time Complexity: O(n log n) — dominated by the sorting step.")
bullet("Guarantees the globally optimal solution for the Fractional Knapsack problem.")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("4.2  Naive Selection Approach (Suboptimal)")
run.font.size  = Pt(12)
run.font.bold  = True
run.font.color.rgb = ACCENT_GOLD
doc.add_paragraph()

bullet("Items are loaded in their original input order — no ratio calculation performed.")
bullet("No sorting or optimization applied; first-come, first-served.")
bullet("Often leaves high-value items behind, resulting in a suboptimal total profit.")
bullet("Serves as a baseline comparison to highlight the advantage of the Greedy approach.")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  5. FEATURES
# ══════════════════════════════════════════════════════════════════════════════
section_heading("5", "Features")
features = [
    "Custom Input — Users can set the knapsack capacity and add items with name, weight, and profit.",
    "Sample Data — One-click load of four predefined cargo items for an instant demo.",
    "Step-by-Step Simulation — Navigate forward and backward through each loading step for both algorithms.",
    "Ship Visualization — A live cargo bar shows the percentage of ship capacity filled at each step.",
    "Real-Time Stats — Displays total profit earned, capacity used, and remaining space at every step.",
    "Profit Comparison Bar Chart — Final side-by-side bar chart comparing total profits of both approaches.",
    "Profit Growth Line Chart — Step-by-step profit accumulation curves for both strategies.",
    "Final Results Panel — Displays the absolute and percentage profit improvement of Greedy over Naive.",
    "Algorithm Toggle — Seamlessly switch between Greedy and Naive simulations mid-session.",
    "Fully Responsive — Works on desktop, tablet, and mobile screens.",
]
for f in features:
    bullet(f)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  6. TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
section_heading("6", "Technology Stack")

tech_table = doc.add_table(rows=8, cols=3)
tech_table.style = 'Table Grid'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Technology", "Version", "Purpose"]
for j, h in enumerate(headers):
    cell = tech_table.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1B4F72")
    run = cell.paragraphs[0].runs[0]
    run.font.bold  = True
    run.font.color.rgb = WHITE
    run.font.size  = Pt(10.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

tech_data = [
    ("React",          "19.x",   "Frontend UI framework"),
    ("TypeScript",     "5.9.x",  "Type-safe JavaScript"),
    ("Vite",           "7.x",    "Dev server & production bundler"),
    ("Tailwind CSS",   "3.4.x",  "Utility-first responsive styling"),
    ("Recharts",       "3.4.x",  "Bar chart & line chart visualizations"),
    ("Lucide React",   "0.554",  "Icon library"),
    ("Vercel",         "—",      "Deployment & hosting"),
]
for i, (tech, ver, purpose) in enumerate(tech_data):
    row = tech_table.rows[i+1]
    row.cells[0].text = tech
    row.cells[1].text = ver
    row.cells[2].text = purpose
    bg = "EBF5FB" if i % 2 == 0 else "FFFFFF"
    for cell in row.cells:
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_TEXT

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  7. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
section_heading("7", "Project Structure")
p = doc.add_paragraph()
run = p.add_run(
    "Fractional Knapsack/\n"
    "├── src/\n"
    "│   ├── FractionalKnapsackDemo.tsx   # Main component — all algorithm logic & UI\n"
    "│   ├── App.tsx                      # Root application entry\n"
    "│   ├── App.css                      # Application-level styles\n"
    "│   ├── index.css                    # Global styles\n"
    "│   └── main.tsx                     # ReactDOM render entry point\n"
    "├── screenshots/                     # App walkthrough screenshots\n"
    "├── index.html                       # HTML entry point\n"
    "├── package.json                     # Project metadata & dependencies\n"
    "├── tailwind.config.js               # Tailwind CSS configuration\n"
    "├── vite.config.ts                   # Vite build configuration\n"
    "└── tsconfig.json                    # TypeScript compiler options"
)
run.font.name  = "Courier New"
run.font.size  = Pt(9.5)
run.font.color.rgb = DARK_TEXT
p.paragraph_format.left_indent = Cm(0.5)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  8. APP SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
section_heading("8", "Application Screenshots")

screenshot_dir = r"d:\Fractional Knapsack\screenshots"
screenshots = [
    ("1_home_setup.png",      "Fig 8.1 — Home: Cargo Loading Setup Page"),
    ("2_sample_data.png",     "Fig 8.2 — Inventory: Sample Data Loaded"),
    ("3_greedy_simulation.png","Fig 8.3 — Simulation: Greedy Algorithm in Action"),
    ("4_naive_simulation.png", "Fig 8.4 — Simulation: Naive Selection Approach"),
    ("5_final_results.png",   "Fig 8.5 — Final Results: Profit Comparison & Charts"),
]

for fname, caption in screenshots:
    fpath = os.path.join(screenshot_dir, fname)
    if os.path.exists(fpath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fpath, width=Inches(5.8))

        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9.5)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x55, 0x6B, 0x82)
        cap.paragraph_format.space_after = Pt(14)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  9. SAMPLE EXECUTION
# ══════════════════════════════════════════════════════════════════════════════
section_heading("9", "Sample Execution Trace")
body(
    "Using the default sample data (Ship capacity = 50 tons):"
)

# Sample data table
sample_table = doc.add_table(rows=5, cols=4)
sample_table.style = 'Table Grid'
sample_table.alignment = WD_TABLE_ALIGNMENT.CENTER

s_headers = ["Item", "Weight (tons)", "Profit (₹L)", "Ratio (₹/ton)"]
for j, h in enumerate(s_headers):
    cell = sample_table.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1B4F72")
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

sample_data = [
    ("Copper Coils",    "10", "60",  "6.0"),
    ("Rice Bags",       "20", "100", "5.0"),
    ("Machinery Parts", "15", "75",  "5.0"),
    ("Cement Blocks",   "30", "120", "4.0"),
]
for i, row_data in enumerate(sample_data):
    row = sample_table.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        bg = "EBF5FB" if i % 2 == 0 else "FFFFFF"
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].runs[0]
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_TEXT
        row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
body("Greedy Algorithm selects items in order of ratio: Copper Coils (6.0) → Rice Bags / Machinery Parts (5.0) → fraction of Cement Blocks.")

p = doc.add_paragraph()
run = p.add_run("  Greedy Total Profit: ₹235 L   |   Naive Total Profit: ₹235 L  →  improvement depends on input order")
run.font.size  = Pt(11)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x1A, 0x8A, 0x44)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  10. LIVE DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
section_heading("10", "Live Deployment")
body(
    "The application is deployed on Vercel and publicly accessible without any "
    "installation or setup. The deployment is continuous — any push to the main "
    "branch automatically triggers a new build and deployment."
)

dep_table = doc.add_table(rows=3, cols=2)
dep_table.style = 'Table Grid'
dep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
dep_info = [
    ("Platform",  "Vercel"),
    ("Live URL",  "https://coding-skills-project-2nd-year.vercel.app/"),
    ("Repository","https://github.com/Bhaumik1904/Coding-Skills-Project-2nd-Year"),
]
for i, (k, v) in enumerate(dep_info):
    bg = "EBF5FB" if i % 2 == 0 else "FFFFFF"
    dep_table.rows[i].cells[0].text = k
    dep_table.rows[i].cells[1].text = v
    for j, cell in enumerate(dep_table.rows[i].cells):
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10.5)
        if j == 0:
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
        else:
            run.font.color.rgb = MID_BLUE
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  11. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
section_heading("11", "Conclusion")
body(
    "The Fractional Knapsack Visualizer successfully achieves its goal of making a "
    "classic algorithm tangible and visually comprehensible. By abstracting the problem "
    "into a cargo-loading scenario, users can intuitively grasp why sorting by profit-to-weight "
    "ratio leads to optimal results."
)
body(
    "The project demonstrates practical skills in React, TypeScript, algorithm design, "
    "data visualization with Recharts, and cloud deployment with Vercel — all core "
    "competencies aligned with the objectives of the Coding Skills course."
)
body(
    "Future enhancements could include a 0/1 Knapsack comparison mode, user-adjustable "
    "animation speed, export-to-PDF for simulation results, and a teacher mode with guided "
    "quiz questions embedded in the simulation."
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  12. TEAM DETAILS
# ══════════════════════════════════════════════════════════════════════════════
section_heading("12", "Team Details")

final_team_table = doc.add_table(rows=6, cols=3)
final_team_table.style = 'Table Grid'
final_team_table.alignment = WD_TABLE_ALIGNMENT.CENTER

f_headers = ["S.No.", "Name", "Enrollment Number"]
for j, h in enumerate(f_headers):
    cell = final_team_table.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1B4F72")
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

final_team = [
    ("1", "Bhaumik Hinunia",        "AP24110010182"),
    ("2", "Manya Srivastava",       "AP24110010171"),
    ("3", "Vardhana Paluvai",       "AP24110012136"),
    ("4", "Jayanti Yadav",          "AP24110010132"),
    ("5", "Vijithendra Nagabhyru",  "AP24110010184"),
]
for i, (sno, name, enroll) in enumerate(final_team):
    row = final_team_table.rows[i+1]
    row.cells[0].text = sno
    row.cells[1].text = name
    row.cells[2].text = enroll
    bg = "EBF5FB" if i % 2 == 0 else "FFFFFF"
    for cell in row.cells:
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Course: Coding Skills  |  Year: 2nd Year  |  Semester: 4th  |  Section: C")
run.font.size  = Pt(11)
run.font.bold  = True
run.font.color.rgb = DARK_BLUE
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"d:\Fractional Knapsack\Fractional_Knapsack_Project_Report.docx"
doc.save(out_path)
print(f"DONE  Report saved -> {out_path}")
